"""
Spracovanie dokumentov s parent-child chunkovaním.
"""

import hashlib
import json
import logging
import re
import subprocess
import sys
from pathlib import Path
from typing import Optional, Literal

import fitz
from langchain_core.documents import Document
from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

ProcessingMode = Literal["fast", "slow", "quality"]

# Bump pri zmene schémy – invaliduje staré cache
_PROCESSOR_CACHE_VERSION = "pc_flat_v1"

# ─── Konfigurácia chunkingu ───────────────────────────────────────────────────
HEADERS_TO_SPLIT = [("#", "h1"), ("##", "h2"), ("###", "h3")]

MIN_PARENT_SIZE  = 400
MAX_PARENT_SIZE  = 2800

CHILD_CHUNK_SIZE    = 600
CHILD_CHUNK_OVERLAP = 100

MIN_CHUNK_CHARS = 60
MARKDOWN_SEPARATORS = ["\n\n", "\n", ". ", "? ", "! ", " ", ""]

PAGE_MARKER_RE = re.compile(r'<!--PAGE:(\d+)-->\n?')

_CONTENT_LIST_SKIP_TYPES = frozenset({
    "header", "footer", "page_number", "aside_text", "page_footnote", "seal",
})


# ─── CUDA detekcia ────────────────────────────────────────────────────────────

def _detect_cuda() -> bool:
    try:
        r = subprocess.run(
            ["nvidia-smi", "--query-gpu=name", "--format=csv,noheader"],
            capture_output=True, text=True, timeout=8,
        )
        if r.returncode == 0 and r.stdout.strip():
            logger.info(f"CUDA GPU: {r.stdout.strip().split(chr(10))[0]}")
            return True
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        pass
    try:
        import torch
        if torch.cuda.is_available():
            logger.info(f"CUDA (torch): {torch.cuda.get_device_name(0)}")
            return True
    except Exception:
        pass
    logger.info("CUDA nedostupná")
    return False


# ─── LaTeX sanitizer ──────────────────────────────────────────────────────────

def sanitize_latex(text: str) -> str:
    def process(latex: str) -> str:
        latex = re.sub(r'\\(?:textstyle|displaystyle|scriptstyle|scriptscriptstyle)\s*', '', latex)
        def unwrap_array(m):
            inner = re.sub(r'^\s*\{[rlc|@{}\s]*\}\s*', '', m.group(1).strip())
            return inner[1:-1].strip() if inner.startswith('{') and inner.endswith('}') else inner
        latex = re.sub(r'\\begin\{array\}(.*?)\\end\{array\}', unwrap_array, latex, flags=re.DOTALL)
        prev = None
        while prev != latex:
            prev = latex
            latex = re.sub(r'\{\s+', '{', latex)
            latex = re.sub(r'\s+\}', '}', latex)
        latex = re.sub(r'(?<![\\a-zA-Z])\{(\\[a-zA-Z]+(?:\[[^\]]*\])?(?:\{[^{}]*\})*)\}', r'\1', latex)
        return re.sub(r'  +', ' ', latex).strip()

    text = re.sub(r'\$\$\s*(.*?)\s*\$\$', lambda m: '$$\n' + process(m.group(1)) + '\n$$', text, flags=re.DOTALL)
    text = re.sub(r'(?<!\$)\$(?!\$)(.*?)(?<!\$)\$(?!\$)', lambda m: '$' + process(m.group(1)) + '$', text)
    return text


# ─── HTML tabuľky → Markdown ──────────────────────────────────────────────────

def _html_tables_to_markdown(text: str) -> str:
    def convert(m):
        rows = re.findall(r'<tr[^>]*>(.*?)</tr>', m.group(0), re.DOTALL | re.IGNORECASE)
        if not rows:
            return m.group(0)
        md = []
        for i, row in enumerate(rows):
            cells = [re.sub(r'<[^>]+>', '', c).strip().replace('\n', ' ')
                     for c in re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', row, re.DOTALL | re.IGNORECASE)]
            if not cells:
                continue
            md.append('| ' + ' | '.join(cells) + ' |')
            if i == 0:
                md.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        return '\n' + '\n'.join(md) + '\n'
    return re.sub(r'<table[^>]*>.*?</table>', convert, text, flags=re.DOTALL | re.IGNORECASE)


# ─── Math bracket normalizácia ────────────────────────────────────────────────

def _normalize_math_brackets(text: str) -> str:
    sqrt_ch = '\u221a'
    pm_ch   = '\u00b1'
    ops     = '\u00d7\u00f7=<>\u2264\u2265\u2260\u2208\u2209\u2282\u222a\u2229~'
    prev = None
    while prev != text:
        prev = text
        text = re.sub(r'\]\s*\[', '', text)
    spat = r'\[' + re.escape(sqrt_ch) + r'\]'
    text = re.sub(spat + r'\s*([0-9]+|[a-zA-Z])\s+', lambda m: f'__SQRT_{m.group(1)}__', text)
    text = re.sub(spat, '__SQRT__', text)
    text = re.sub(r'(?<=[a-zA-Z0-9\)]) ?\[([^\[\]\n]{1,50})\]',
                  lambda m: '[' + m.group(1) + ']' if (',' in m.group(1) and len(m.group(1)) > 8)
                  else '^{' + m.group(1) + '}', text)
    text = text.replace('\u0338=', '\u2260').replace('\u0338', '')
    text = re.sub(r'\[\s*([' + r'+\-' + pm_ch + ops + r'])\s*\]',
                  lambda m: ' ' + m.group(1) + ' ', text)
    text = re.sub(r'\[\s*\]', ' ', text)
    text = re.sub(r'__SQRT_([^_]+)__', lambda m: r'\sqrt[' + m.group(1) + ']', text)
    text = re.sub(r'__SQRT__', lambda m: r'\sqrt', text)
    pm_cls = r'+\-' + pm_ch
    text = re.sub(r'\^\{([^}]+)\}',
                  lambda m: '^{' + re.sub(r'\s*([' + pm_cls + r'])\s*',
                                          lambda n: ' ' + n.group(1) + ' ', m.group(1)).strip() + '}',
                  text)
    return re.sub(r' {2,}', ' ', text).strip()


# ─── Tabuľky – shadow text pre lepší retrieval ───────────────────────────────

def _expand_tables_for_retrieval(text: str) -> str:
    def table_to_text(match: re.Match) -> str:
        table = match.group(0)
        lines = [l.strip() for l in table.split('\n') if l.strip()]
        rows  = [l for l in lines if not re.match(r'^\|[-| :]+\|$', l)]
        if len(rows) < 2:
            return table
        headers   = [c.strip() for c in rows[0].strip('|').split('|')]
        text_rows = []
        for row in rows[1:]:
            cells = [c.strip() for c in row.strip('|').split('|')]
            if len(cells) == len(headers):
                parts = [f"{h}: {c}" for h, c in zip(headers, cells) if c]
                if parts:
                    text_rows.append(", ".join(parts))
        if text_rows:
            return f"{table}\n" + "\n".join(text_rows)
        return table

    return re.sub(r'(\|.+\|\n)+', table_to_text, text, flags=re.MULTILINE)


# ─── Parent-child chunker ─────────────────────────────────────────────────────

class ParentChildChunker:
    """
    Rozdelí markdown na parent chunky (sekcie) a child chunky (kusy sekcií).
    Parent chunky obsahujú úplný kontext sekcie pre LLM.
    Child chunky sú malé kusy indexované vo FAISS.

    Každý chunk nesie: source, page, page_start, page_end, pages, parent_id,
    h1/h2/h3, heading_path (odvodené z h1-h3).
    """

    def __init__(self):
        self._parent_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=HEADERS_TO_SPLIT,
            strip_headers=False,
        )
        self._child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHILD_CHUNK_SIZE,
            chunk_overlap=CHILD_CHUNK_OVERLAP,
            separators=MARKDOWN_SEPARATORS,
        )

    def chunk(
        self,
        md_text: str,
        source: str,
        base_page: int = 1,
    ) -> tuple[list[Document], list[Document]]:
        """
        Vráti (parents, children). Parent chunky majú parent_id v metadata,
        child chunky tiež (ukazujú späť na rodiča).
        """

        page_texts = self._build_full_page_texts(md_text)
        pages      = self._split_pages(md_text)
        has_markers = bool(pages and not (len(pages) == 1 and pages[0][0] == 1))

        all_parent_candidates: list[Document] = []

        for page_no, page_text in pages:
            page_text_expanded = _expand_tables_for_retrieval(page_text)
            page_docs = self._parent_splitter.split_text(page_text_expanded)

            for d in page_docs:
                d.metadata.setdefault("source", source)
                d.metadata["page"]       = page_no
                d.metadata["page_start"] = page_no
                d.metadata["page_end"]   = page_no
                d.metadata["pages"]      = [page_no]

            all_parent_candidates.extend(page_docs)

        merged = self._merge_small(all_parent_candidates)
        split  = self._split_large(merged)

        if not has_markers and page_texts:
            for doc in split:
                p_start, p_end, p_pages, primary = self._detect_page_span_from_fulltext(
                    doc.page_content, page_texts, base_page
                )
                doc.metadata["page"]       = primary
                doc.metadata["page_start"] = p_start
                doc.metadata["page_end"]   = p_end
                doc.metadata["pages"]      = p_pages

        cleaned = self._clean_small(split)

        stem = Path(source).stem
        parents: list[Document] = []
        child_docs: list[Document] = []

        for i, p_chunk in enumerate(cleaned):
            parent_id = f"{stem}_parent_{i}"
            p_chunk.metadata.setdefault("source", source)
            p_chunk.metadata["parent_id"] = parent_id

            # Odvoď heading_path z h1/h2/h3 (pre RAGGraph compress node)
            heading_parts = [
                p_chunk.metadata[k] for k in ("h1", "h2", "h3")
                if p_chunk.metadata.get(k)
            ]
            if heading_parts:
                p_chunk.metadata["heading_path"] = " › ".join(heading_parts)

            parents.append(p_chunk)

            local_parents = self._split_parent_by_pages(p_chunk, page_texts)

            for local_parent in local_parents:
                lp_page       = local_parent.metadata.get("page",       p_chunk.metadata.get("page", base_page))
                lp_page_start = local_parent.metadata.get("page_start", lp_page)
                lp_page_end   = local_parent.metadata.get("page_end",   lp_page)
                lp_pages      = local_parent.metadata.get("pages",      [lp_page])

                children = self._child_splitter.split_documents([local_parent])

                for child in children:
                    if len(child.page_content.strip()) < MIN_CHUNK_CHARS:
                        continue

                    child.metadata["parent_id"]  = parent_id
                    child.metadata.setdefault("source", source)
                    child.metadata["page"]       = lp_page
                    child.metadata["page_start"] = lp_page_start
                    child.metadata["page_end"]   = lp_page_end
                    child.metadata["pages"]      = lp_pages

                    for k in ("h1", "h2", "h3"):
                        if k in p_chunk.metadata and k not in child.metadata:
                            child.metadata[k] = p_chunk.metadata[k]

                    if p_chunk.metadata.get("heading_path"):
                        child.metadata["heading_path"] = p_chunk.metadata["heading_path"]

                    child_docs.append(child)

        logger.info(
            f"chunk(): {len(parents)} parents, {len(child_docs)} children"
            + (f", markery aktívne" if has_markers else f" (bez PAGE markerov – base_page={base_page})")
        )

        return parents, child_docs

    # ── Page parsing ──────────────────────────────────────────────────────────

    def _split_pages(self, md_text: str) -> list[tuple[int, str]]:
        sections = PAGE_MARKER_RE.split(md_text)
        pages = []

        i = 1
        while i < len(sections) - 1:
            try:
                page_no = int(sections[i])
                content = sections[i + 1].strip()
                if content:
                    pages.append((page_no, content))
            except (ValueError, IndexError):
                pass
            i += 2

        if not pages and md_text.strip():
            content = PAGE_MARKER_RE.sub('', md_text).strip()
            pages.append((1, content))

        return pages

    def _build_full_page_texts(self, md_text: str) -> dict[int, str]:
        pages = self._split_pages(md_text)
        return {page_no: text.lower() for page_no, text in pages}

    def _build_page_map(self, md_text: str) -> dict[str, int]:
        page_map: dict[str, int] = {}
        sections = PAGE_MARKER_RE.split(md_text)

        i = 1
        while i < len(sections) - 1:
            try:
                page_no = int(sections[i])
                content = sections[i + 1] if i + 1 < len(sections) else ""
                for line in content.split('\n'):
                    line_s = line.strip()
                    if len(line_s) >= 15:
                        clean = line_s.lstrip('#').strip()
                        if len(clean) >= 8:
                            page_map[clean[:80].lower()] = page_no
                        if not line_s.startswith('#'):
                            page_map[line_s[:80].lower()] = page_no
            except (ValueError, IndexError):
                pass
            i += 2

        return page_map

    def _detect_page(self, content: str, page_map: dict[str, int], fallback: int) -> int:
        if not page_map:
            return fallback
        content_lower = content.lower()
        votes: dict[int, int] = {}
        for snippet, page_no in page_map.items():
            if snippet in content_lower:
                votes[page_no] = votes.get(page_no, 0) + 1
        if not votes:
            return fallback
        return max(votes, key=lambda p: votes[p])

    def _detect_page_span_from_fulltext(
        self,
        content: str,
        page_texts: dict[int, str],
        fallback: int,
    ) -> tuple[int, int, list[int], int]:
        if not page_texts:
            return fallback, fallback, [fallback], fallback

        c = content.lower().strip()
        if not c:
            return fallback, fallback, [fallback], fallback

        probes = []
        for line in c.split("\n"):
            line = line.strip()
            if len(line) >= 15:
                probes.append(line[:120])

        if not probes:
            probes = [c[:160]]

        scores = {}
        for page_no, page_text in page_texts.items():
            score = 0
            for probe in probes:
                if probe in page_text:
                    score += 1
            if score > 0:
                scores[page_no] = score

        if not scores:
            from difflib import SequenceMatcher
            for page_no, page_text in page_texts.items():
                ratio = SequenceMatcher(None, c[:200], page_text[:500]).ratio()
                if ratio > 0.3:
                    scores[page_no] = ratio

        if not scores:
            guessed = self._detect_page(
                content,
                self._build_page_map_from_page_texts(page_texts),
                fallback,
            )
            return guessed, guessed, [guessed], guessed

        matched_pages = sorted(scores.keys())

        if not matched_pages:
            return fallback, fallback, [fallback], fallback

        primary_page = max(scores, key=scores.get)
        return matched_pages[0], matched_pages[-1], matched_pages, primary_page

    def _build_page_map_from_page_texts(self, page_texts: dict[int, str]) -> dict[str, int]:
        page_map = {}
        for page_no, text in page_texts.items():
            for line in text.split("\n"):
                line = line.strip()
                if len(line) >= 20:
                    page_map[line[:80].lower()] = page_no
        return page_map

    def _split_parent_by_pages(self, parent_doc: Document, page_texts: dict[int, str]) -> list[Document]:
        pages = parent_doc.metadata.get("pages", [])
        if len(pages) <= 1:
            return [parent_doc]

        content = parent_doc.page_content
        subdocs = []

        for pg in pages:
            page_text = page_texts.get(pg, "")
            if not page_text:
                continue

            matched_lines = []
            for line in content.split("\n"):
                ls = line.strip()
                if len(ls) >= 20 and ls.lower() in page_text:
                    matched_lines.append(line)

            if matched_lines:
                subdocs.append(Document(
                    page_content="\n".join(matched_lines),
                    metadata={
                        **parent_doc.metadata,
                        "page": pg,
                        "page_start": pg,
                        "page_end": pg,
                        "pages": [pg],
                    }
                ))

        return subdocs if subdocs else [parent_doc]

    # ── Merge / split / clean ─────────────────────────────────────────────────

    def _merge_metadata(self, a: dict, b: dict) -> dict:
        merged = dict(a)
        merged["source"] = a.get("source", b.get("source"))

        a_pages = a.get("pages", [])
        b_pages = b.get("pages", [])
        pages = sorted(set(a_pages + b_pages))
        if pages:
            merged["pages"] = pages
            merged["page_start"] = min(pages)
            merged["page_end"] = max(pages)
            merged["page"] = merged["page_start"]

        for key in ["h1", "h2", "h3"]:
            if b.get(key):
                merged[key] = b[key]
            elif a.get(key):
                merged[key] = a[key]

        for k, v in b.items():
            if k not in merged:
                merged[k] = v

        return merged

    def _merge_small(self, chunks: list, page_map: dict[str, int] = None) -> list:
        if not chunks:
            return []

        merged = []
        current = None

        for chunk in chunks:
            if current is None:
                current = chunk
                continue

            same_page_span = (
                current.metadata.get("page_start") == chunk.metadata.get("page_start")
                and current.metadata.get("page_end") == chunk.metadata.get("page_end")
            )

            if not same_page_span:
                merged.append(current)
                current = chunk
                continue

            current.page_content += "\n\n" + chunk.page_content
            current.metadata = self._merge_metadata(current.metadata, chunk.metadata)

            if len(current.page_content) >= MIN_PARENT_SIZE:
                merged.append(current)
                current = None

        if current is not None:
            if merged:
                last_same_page = (
                    merged[-1].metadata.get("page_start") == current.metadata.get("page_start")
                    and merged[-1].metadata.get("page_end") == current.metadata.get("page_end")
                )
                if last_same_page:
                    merged[-1].page_content += "\n\n" + current.page_content
                    merged[-1].metadata = self._merge_metadata(merged[-1].metadata, current.metadata)
                else:
                    merged.append(current)
            else:
                merged.append(current)

        return merged

    def _split_large(self, chunks: list) -> list:
        result   = []
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=MAX_PARENT_SIZE,
            chunk_overlap=CHILD_CHUNK_OVERLAP,
            separators=MARKDOWN_SEPARATORS,
        )

        for chunk in chunks:
            if len(chunk.page_content) <= MAX_PARENT_SIZE:
                result.append(chunk)
            else:
                split_docs = splitter.split_documents([chunk])
                for d in split_docs:
                    d.metadata = dict(chunk.metadata)
                result.extend(split_docs)

        return result

    def _clean_small(self, chunks: list) -> list:
        cleaned = []

        for i, chunk in enumerate(chunks):
            if len(chunk.page_content.strip()) < MIN_CHUNK_CHARS:
                if cleaned:
                    last_same_page = (
                        cleaned[-1].metadata.get("page_start") == chunk.metadata.get("page_start")
                        and cleaned[-1].metadata.get("page_end") == chunk.metadata.get("page_end")
                    )
                    if last_same_page:
                        cleaned[-1].page_content += "\n\n" + chunk.page_content
                        cleaned[-1].metadata = self._merge_metadata(cleaned[-1].metadata, chunk.metadata)
                    else:
                        cleaned.append(chunk)
                elif i < len(chunks) - 1:
                    next_same_page = (
                        chunks[i+1].metadata.get("page_start") == chunk.metadata.get("page_start")
                        and chunks[i+1].metadata.get("page_end") == chunk.metadata.get("page_end")
                    )
                    if next_same_page:
                        chunks[i+1].page_content = chunk.page_content + "\n\n" + chunks[i+1].page_content
                        chunks[i+1].metadata = self._merge_metadata(chunk.metadata, chunks[i+1].metadata)
                    else:
                        cleaned.append(chunk)
                else:
                    cleaned.append(chunk)
            else:
                cleaned.append(chunk)

        return cleaned


# ─── Hlavná trieda ────────────────────────────────────────────────────────────

class DocumentProcessor:

    def __init__(self, notebook_path: Path, chunk_size: int = CHILD_CHUNK_SIZE,
                 chunk_overlap: int = CHILD_CHUNK_OVERLAP):
        self.notebook_path = notebook_path
        self.files_dir     = notebook_path / "files"
        self.images_dir    = notebook_path / "page_images"
        self._cache_dir    = notebook_path / ".page_cache"
        self._cache_dir.mkdir(parents=True, exist_ok=True)
        self._chunker      = ParentChildChunker()
        self._cuda: Optional[bool] = None

    def process_file(
        self,
        filename: str,
        mode: ProcessingMode = "fast",
    ) -> tuple[list[Document], list[Document]]:
        # Normalizuj "quality" → "slow" (app.py niekde posiela quality)
        if mode == "quality":
            mode = "slow"

        file_path = self.files_dir / filename
        suffix    = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._process_pdf(file_path, filename, mode)
        elif suffix in (".xlsx", ".xls"):
            return self._process_xlsx(file_path, filename)
        elif suffix in (".docx", ".doc"):
            return self._process_docx(file_path, filename)
        logger.warning(f"Nepodporovaný formát: {suffix}")
        return [], []

    # ─── PDF ──────────────────────────────────────────────────────────────────

    def _process_pdf(self, file_path: Path, filename: str, mode: ProcessingMode):
        doc_hash = self._doc_hash(file_path, mode)

        cached = self._load_cache(doc_hash)
        if cached is not None:
            parents, children = cached
            if children:
                pages = [c.metadata.get("page", 0) for c in children[:min(6, len(children))]]
                if len(pages) > 2 and len(set(pages)) == 1:
                    logger.info(f"Cache invalidovaná (zlé page metadata): {filename}")
                    cached = None

        if cached is not None:
            logger.info(f"Cache hit: '{filename}' → {len(cached[0])} parents, {len(cached[1])} children")
            self._render_pdf_pages(file_path, filename)
            return cached

        md_text = (
            self._extract_fast(file_path) if mode == "fast"
            else self._extract_slow(file_path)
        )

        parents, children = self._chunker.chunk(md_text, filename)

        if children:
            try:
                from langdetect import detect
                sample = " ".join(c.page_content[:100] for c in children[:5])
                doc_language = detect(sample) if sample.strip() else "en"
            except Exception:
                doc_language = "en"

            for child in children:
                child.metadata["doc_language"] = doc_language

        self._save_cache(doc_hash, (parents, children))
        self._render_pdf_pages(file_path, filename)

        logger.info(
            f"'{filename}': {len(parents)} parents, {len(children)} children (mode={mode})"
        )
        return parents, children

    # ─── Fast: pymupdf4llm ────────────────────────────────────────────────────

    def _extract_fast(self, file_path: Path) -> str:
        """Rýchla extrakcia textu z PDF cez pymupdf4llm."""
        import pymupdf4llm

        pages = pymupdf4llm.to_markdown(str(file_path), page_chunks=True)

        parts = []
        for i, page_data in enumerate(pages):
            text = page_data.get("text", "").strip()
            if not text:
                continue
            page_number = i + 1  # Strana v PDF (1-based indexing)
            parts.append(
                f"<!--PAGE:{page_number}-->\n"
                f"{text}"
            )
        return "\n\n".join(parts)

    # ─── Slow: MinerU ────────────────────────────────────────────────────────

    def _extract_slow(self, file_path: Path) -> str:
        """Kvalitná extrakcia z PDF cez MinerU. Spracováva výstup z content_list.json."""
        if self._cuda is None:
            self._cuda = _detect_cuda()

        output_dir = file_path.parent / ".mineru_cache" / file_path.stem
        output_dir.mkdir(parents=True, exist_ok=True)

        # Cache hit – content_list.json už existuje
        content_list_path = self._find_mineru_content_list(output_dir, file_path.stem)
        if content_list_path and content_list_path.exists():
            logger.info(f"MinerU cache HIT: {content_list_path.name}")
            md = self._build_md_from_content_list(content_list_path)
            if md:
                return md

        # Spustenie MinerU
        cmd = [
            sys.executable, "-m", "mineru.cli.client",
            "-p", str(file_path),
            "-o", str(output_dir),
            "--backend", "pipeline",
        ]
        if self._cuda:
            cmd += ["--device", "cuda"]
        logger.info(f"Slow (MinerU{'+ CUDA' if self._cuda else ' CPU'}): '{file_path.name}'...")

        try:
            result = subprocess.run(cmd, capture_output=True, text=True)
        except FileNotFoundError:
            logger.error('MinerU nie je nainštalovaný. Spusti: pip install "mineru[pipeline]"')
            return ""
        except Exception as e:
            logger.error(f"MinerU: {e}")
            return ""

        if result.returncode != 0:
            logger.error(f"MinerU exit {result.returncode}: {result.stderr[-300:]}")
            return ""

        # Spracovanie výstupu cez content_list.json
        content_list_path = self._find_mineru_content_list(output_dir, file_path.stem)
        if not content_list_path:
            logger.error("MinerU neprodukoval content_list.json")
            return ""

        logger.info(f"Používam content_list.json: {content_list_path.name}")
        return self._build_md_from_content_list(content_list_path)

    # ─── content_list.json lokácia ────────────────────────────────────────────

    def _find_mineru_content_list(self, output_dir: Path, stem: str) -> Optional[Path]:
        candidates = [
            output_dir / stem / "auto" / f"{stem}_content_list.json",
            output_dir / "auto" / f"{stem}_content_list.json",
            output_dir / stem / f"{stem}_content_list.json",
            output_dir / f"{stem}_content_list.json",
        ]
        for p in candidates:
            if p.exists():
                return p

        found = sorted(output_dir.rglob("*content_list*.json"))
        non_v2 = [f for f in found if "v2" not in f.name]
        if non_v2:
            return non_v2[0]
        return found[0] if found else None

    # ─── Markdown z content_list.json ─────────────────────────────────────────

    def _build_md_from_content_list(self, content_list_path: Path) -> str:
        try:
            raw = json.loads(content_list_path.read_text(encoding="utf-8"))
        except Exception as e:
            logger.error(f"content_list.json load: {e}")
            return ""

        blocks: list[dict] = []
        if raw and isinstance(raw[0], list):
            for page_idx, page_blocks in enumerate(raw):
                for b in page_blocks:
                    b.setdefault("page_idx", page_idx)
                    blocks.append(b)
        else:
            blocks = raw

        if not blocks:
            logger.warning("content_list.json je prázdny")
            return ""

        result:       list[str]    = []
        current_page: Optional[int] = None

        for block in blocks:
            try:
                physical_page = int(block.get("page_idx", 0)) + 1
            except Exception:
                physical_page = 1

            physical_page = max(1, physical_page)
            block_type = block.get("type", "text")

            if block_type in _CONTENT_LIST_SKIP_TYPES:
                continue

            if physical_page != current_page:
                result.append(f"<!--PAGE:{physical_page}-->")
                current_page = physical_page

            if block_type == "text":
                text = block.get("text", "").strip()
                if not text:
                    continue
                level  = block.get("text_level", 0)
                prefix = {1: "# ", 2: "## ", 3: "### "}.get(level, "")
                result.append(prefix + text)

            elif block_type == "equation":
                text = block.get("text", "").strip()
                if not text:
                    continue
                if not text.startswith("$$"):
                    text = f"$$\n{text}\n$$"
                result.append(text)

            elif block_type == "table":
                captions  = block.get("table_caption", [])
                footnotes = block.get("table_footnote", [])
                body      = block.get("table_body", "")

                if captions:
                    result.append(f"*{captions[0].strip()}*")
                if body:
                    result.append(_html_tables_to_markdown(body))
                if footnotes:
                    result.append(f"_{footnotes[0].strip()}_")

            elif block_type in ("image", "chart"):
                captions = (
                    block.get("image_caption")
                    or block.get("chart_caption")
                    or []
                )
                if captions and captions[0].strip():
                    result.append(f"[Obrázok: {captions[0].strip()[:120]}]")
                else:
                    result.append("[Obrázok]")

            elif block_type == "list":
                items = block.get("list_items", [])
                for item in items:
                    s = item.strip()
                    if s:
                        result.append(f"- {s}")

            elif block_type == "code":
                captions = block.get("code_caption", [])
                if captions:
                    result.append(f"*{captions[0].strip()}*")
                body = block.get("code_body", "").strip()
                if body:
                    result.append(f"```\n{body}\n```")

            else:
                text = block.get("text", "").strip()
                if text:
                    result.append(text)

        if not result:
            return ""

        parts: list[str] = []
        for line in result:
            if line.startswith("<!--PAGE:"):
                parts.append(line)
            else:
                parts.append(line)
                parts.append("")

        md = "\n".join(parts).strip()
        md = sanitize_latex(md)
        md = _normalize_math_brackets(md)

        page_count = len(set(
            m.group(1) for m in PAGE_MARKER_RE.finditer(md)
        ))
        logger.info(
            f"_build_md_from_content_list: {len(blocks)} blokov, "
            f"{page_count} unique strán, {len(md)} znakov"
        )
        return md

    # ─── Cache ────────────────────────────────────────────────────────────────

    def _doc_hash(self, fp: Path, mode: str = "fast") -> str:
        """Hash zahŕňa mód – prepnutie fast/slow invaliduje cache."""
        try:
            s = fp.stat()
            return hashlib.sha256(
                f"{_PROCESSOR_CACHE_VERSION}|{mode}|{fp.name}_{s.st_size}_{s.st_mtime}".encode()
            ).hexdigest()[:16]
        except Exception:
            return hashlib.sha256(
                f"{_PROCESSOR_CACHE_VERSION}|{mode}|{str(fp)}".encode()
            ).hexdigest()[:16]

    def _load_cache(self, h: str):
        f = self._cache_dir / f"pc_{h}.json"
        if not f.exists():
            return None
        try:
            data     = json.loads(f.read_text(encoding="utf-8"))
            parents  = [
                Document(page_content=item["content"], metadata=item["meta"])
                for item in data["parents"]
            ]
            children = [
                Document(page_content=item["content"], metadata=item["meta"])
                for item in data["children"]
            ]
            return parents, children
        except Exception as e:
            logger.warning(f"Cache load: {e}")
            return None

    def _save_cache(self, h: str, result: tuple) -> None:
        parents, children = result
        try:
            data = {
                "parents":  [
                    {"content": d.page_content, "meta": d.metadata}
                    for d in parents
                ],
                "children": [
                    {"content": d.page_content, "meta": d.metadata}
                    for d in children
                ],
            }
            (self._cache_dir / f"pc_{h}.json").write_text(
                json.dumps(data, ensure_ascii=False),
                encoding="utf-8",
            )
        except Exception as e:
            logger.warning(f"Cache save: {e}")

    # ─── Renderovanie PDF stránok ─────────────────────────────────────────────

    def _render_pdf_pages(self, file_path: Path, filename: str) -> None:
        img_dir = self.images_dir / filename
        img_dir.mkdir(parents=True, exist_ok=True)
        try:
            pdf = fitz.open(str(file_path))
            for i in range(len(pdf)):
                p = img_dir / f"page_{i+1}.png"
                if p.exists():
                    continue
                pdf[i].get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False).save(str(p))
            pdf.close()
        except Exception as e:
            logger.error(f"Renderovanie: {e}")

    # ─── XLSX ─────────────────────────────────────────────────────────────────

    def _process_xlsx(self, file_path: Path, filename: str):
        try:
            import pandas as pd
            xl  = pd.ExcelFile(str(file_path))
            md  = "\n\n".join(
                f"## {s}\n\n{xl.parse(s).to_string(index=False)}"
                for s in xl.sheet_names
            )
            return self._chunker.chunk(md, filename)
        except Exception as e:
            logger.error(f"XLSX: {e}", exc_info=True)
            return [], []

    # ─── DOCX ─────────────────────────────────────────────────────────────────

    def _process_docx(self, file_path: Path, filename: str):
        try:
            from docx import Document as DocxDoc
            doc   = DocxDoc(str(file_path))
            lines = []
            for p in doc.paragraphs:
                if not p.text.strip():
                    lines.append("")
                    continue
                s = p.style.name.lower() if p.style else ""
                prefix = (
                    "# "   if "heading 1" in s else
                    "## "  if "heading 2" in s else
                    "### " if "heading 3" in s else
                    ""
                )
                lines.append(prefix + p.text)
            return self._chunker.chunk("\n\n".join(lines), filename)
        except Exception as e:
            logger.error(f"DOCX: {e}", exc_info=True)
            return [], []

    def get_page_image_path(self, filename: str, page_no: int) -> Optional[Path]:
        p = self.images_dir / filename / f"page_{page_no}.png"
        return p if p.exists() else None